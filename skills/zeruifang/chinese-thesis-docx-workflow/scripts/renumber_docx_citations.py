#!/usr/bin/env python3
"""Renumber bracket citations and bibliography by first appearance."""
from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document

CITE_RE = re.compile(r"\[([0-9][0-9,\-\s]*?)\]")
REF_RE = re.compile(r"^\[(\d+)\]\s*(.*)")


def expand_nums(text: str) -> list[int]:
    nums: list[int] = []
    for part in re.split(r"\s*,\s*", text.strip()):
        if not part:
            continue
        if "-" in part:
            a_s, b_s = part.split("-", 1)
            a, b = int(a_s), int(b_s)
            step = 1 if b >= a else -1
            nums.extend(range(a, b + step, step))
        else:
            nums.append(int(part))
    return nums


def compress_nums(nums: list[int]) -> str:
    out: list[str] = []
    i = 0
    while i < len(nums):
        start = nums[i]
        j = i
        while j + 1 < len(nums) and nums[j + 1] == nums[j] + 1:
            j += 1
        if j - i + 1 >= 3:
            out.append(f"{start}-{nums[j]}")
        else:
            out.extend(str(nums[k]) for k in range(i, j + 1))
        i = j + 1
    return ",".join(out)


def find_reference_start(doc: Document) -> int:
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip() == "参考文献":
            return i
    raise SystemExit("reference heading '参考文献' not found")


def analyze(doc: Document) -> tuple[int, dict[int, str], list[int], list[int], list[tuple[int, int, str]]]:
    ref_start = find_reference_start(doc)
    refs: dict[int, str] = {}
    ref_paras: list[int] = []
    for i in range(ref_start + 1, len(doc.paragraphs)):
        text = doc.paragraphs[i].text.strip()
        match = REF_RE.match(text)
        if match:
            refs[int(match.group(1))] = match.group(2)
            ref_paras.append(i)

    first_order: list[int] = []
    missing: list[tuple[int, int, str]] = []
    for i in range(ref_start):
        text = doc.paragraphs[i].text
        for match in CITE_RE.finditer(text):
            for old in expand_nums(match.group(1)):
                if old not in refs:
                    missing.append((i, old, text[:120]))
                elif old not in first_order:
                    first_order.append(old)
    return ref_start, refs, ref_paras, first_order, missing


def validate(doc: Document) -> bool:
    ref_start, refs, _ref_paras, first_order, missing = analyze(doc)
    if missing:
        print("missing citations:")
        for item in missing:
            print(item)
        return False
    expected = list(range(1, len(refs) + 1))
    ref_nums = []
    for i in range(ref_start + 1, len(doc.paragraphs)):
        match = REF_RE.match(doc.paragraphs[i].text.strip())
        if match:
            ref_nums.append(int(match.group(1)))
    print("body_first_order_ok", first_order == expected)
    print("reference_sequence_ok", ref_nums == expected)
    print("citation_count", len(first_order), "reference_count", len(ref_nums))
    return first_order == expected and ref_nums == expected


def renumber(doc: Document, remove_uncited: bool) -> None:
    ref_start, refs, ref_paras, first_order, missing = analyze(doc)
    if missing:
        raise SystemExit(f"missing citations: {missing[:10]}")
    uncited = [n for n in sorted(refs) if n not in first_order]
    ordered = first_order if remove_uncited else first_order + uncited
    old_to_new = {old: i + 1 for i, old in enumerate(ordered)}

    def replace_text(text: str) -> str:
        def sub(match: re.Match[str]) -> str:
            new_nums: list[int] = []
            for old in expand_nums(match.group(1)):
                new = old_to_new.get(old)
                if new is not None and new not in new_nums:
                    new_nums.append(new)
            return "[" + compress_nums(new_nums) + "]"

        return CITE_RE.sub(sub, text)

    changed = 0
    for i in range(ref_start):
        para = doc.paragraphs[i]
        new_text = replace_text(para.text)
        if new_text != para.text:
            para.clear()
            para.add_run(new_text)
            changed += 1

    for offset, old_num in enumerate(ordered):
        para = doc.paragraphs[ref_paras[offset]]
        para.clear()
        para.add_run(f"[{offset + 1}] {refs[old_num]}")
    for para_i in ref_paras[len(ordered):]:
        doc.paragraphs[para_i].clear()

    print("renumbered_body_paragraphs", changed)
    print("references_kept", len(ordered), "uncited", uncited)


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("input", type=Path)
    parser.add_argument("--output", type=Path)
    parser.add_argument("--check-only", action="store_true")
    parser.add_argument("--remove-uncited", action="store_true")
    args = parser.parse_args()

    doc = Document(args.input)
    if args.check_only:
        return 0 if validate(doc) else 1
    if not args.output:
        raise SystemExit("--output is required unless --check-only is used")
    renumber(doc, args.remove_uncited)
    doc.save(args.output)
    print("saved", args.output)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
