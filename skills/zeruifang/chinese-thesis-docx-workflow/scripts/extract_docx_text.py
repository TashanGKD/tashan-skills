#!/usr/bin/env python3
"""Extract paragraph text and structure from a DOCX."""
from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("docx", type=Path)
    parser.add_argument("--start", type=int, default=0)
    parser.add_argument("--end", type=int)
    parser.add_argument("--grep", help="Only print paragraphs containing this text")
    args = parser.parse_args()

    doc = Document(args.docx)
    end = args.end if args.end is not None else len(doc.paragraphs)
    for i, para in enumerate(doc.paragraphs[args.start:end], start=args.start):
        text = para.text.strip()
        has_drawing = bool(para._p.xpath(".//w:drawing"))
        if args.grep and args.grep not in text:
            continue
        if text or has_drawing:
            print(f"{i}: drawing={has_drawing} style={para.style.name} text={text}")
    print(f"paragraphs={len(doc.paragraphs)} inline_shapes={len(doc.inline_shapes)}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
